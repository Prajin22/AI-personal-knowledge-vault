import os
import io
import re
from typing import Dict, List, Optional, Union, BinaryIO
from pathlib import Path

# Optional heavy dependencies are imported lazily to avoid crash at import time
_PYPDF2 = None
_DOCX = None
_PANDAS = None
_OPENPYXL = None

try:
    import PyPDF2 as _PYPDF2
except Exception:
    _PYPDF2 = None

try:
    import docx as _DOCX
except Exception:
    _DOCX = None

try:
    import pandas as _PANDAS
except Exception:
    _PANDAS = None

try:
    import openpyxl as _OPENPYXL
except Exception:
    _OPENPYXL = None

class DocumentProcessor:
    """Process various document types (PDF, DOCX, XLSX) and extract text content."""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf_text,
            '.docx': self._extract_docx_text,
            '.xlsx': self._extract_excel_text,
            '.xls': self._extract_excel_text,
            '.csv': self._extract_csv_text
        }
    
    def process_document(self, file_path: Union[str, BinaryIO], file_extension: str) -> str:
        """Process a document and return its text content."""
        file_extension = file_extension.lower()
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        extractor = self.supported_formats[file_extension]
        text = extractor(file_path)

        processed_text = self._process_qa_content(text)
        return processed_text if processed_text else text
    
    def _extract_pdf_text(self, file_path: Union[str, BinaryIO]) -> str:
        """Extract text from PDF file."""
        text = []
        try:
            if hasattr(file_path, 'read'):
                # Handle file-like object - read to bytes first
                file_bytes = file_path.read()
                pdf_reader = _PYPDF2.PdfReader(io.BytesIO(file_bytes))
            else:
                # Handle file path
                with open(file_path, 'rb') as f:
                    pdf_reader = _PYPDF2.PdfReader(f)
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            
            return '\n'.join(text)
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
        finally:
            # Reset file pointer if it's a file-like object
            if hasattr(file_path, 'seek'):
                try:
                    file_path.seek(0)
                except:
                    pass
    
    def _extract_docx_text(self, file_path: Union[str, BinaryIO]) -> str:
        """Extract text from DOCX file."""
        try:
            if hasattr(file_path, 'read'):
                # Handle file-like object
                doc = _DOCX.Document(io.BytesIO(file_path.read()))
            else:
                # Handle file path
                doc = _DOCX.Document(file_path)
            
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def _extract_excel_text(self, file_path: Union[str, BinaryIO]) -> str:
        """Extract text from Excel file (XLSX/XLS)."""
        try:
            if hasattr(file_path, 'read'):
                # Handle file-like object
                df = _PANDAS.read_excel(io.BytesIO(file_path.read()))
            else:
                # Handle file path
                df = _PANDAS.read_excel(file_path)
            
            # Convert all cells to strings and join with newlines
            return '\n'.join(df.astype(str).values.flatten().tolist())
        except Exception as e:
            raise Exception(f"Error extracting text from Excel: {str(e)}")
    
    def _extract_csv_text(self, file_path: Union[str, BinaryIO]) -> str:
        """Extract text from CSV file."""
        try:
            if hasattr(file_path, 'read'):
                # Handle file-like object
                df = _PANDAS.read_csv(io.StringIO(file_path.read().decode('utf-8')))
            else:
                # Handle file path
                df = _PANDAS.read_csv(file_path)
            
            # Convert all cells to strings and join with newlines
            return '\n'.join(df.astype(str).values.flatten().tolist())
        except Exception as e:
            raise Exception(f"Error extracting text from CSV: {str(e)}")
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.supported_formats.keys())
    
    def is_supported(self, file_extension: str) -> bool:
        """Check if a file extension is supported."""
        return file_extension.lower() in self.supported_formats

    def _extract_qa_pairs(self, text: str) -> List[Dict[str, str]]:
        if not text:
            return []

        pairs: List[Dict[str, str]] = []

        md_matches = re.findall(
            r"^##\s*(.+?)\s*\n+(.+?)(?=^##\s*|\Z)",
            text,
            re.DOTALL | re.MULTILINE,
        )
        for q, a in md_matches:
            q = q.strip()
            a = re.sub(r"\s+", " ", a).strip()
            if q and a:
                pairs.append({"question": q, "answer": a})
        if pairs:
            return pairs

        question_patterns = [
            r"(?:Question|Q)[:\s]\s*(.+?)\s*(?:\n|$)(.+?)(?=(?:\n\s*\n)|(?:Question|Q)[:\s]|\Z)",
            r"(\d+\.\s*.+?\?)\s*(.+?)(?=(?:\n\s*\n)|(?:\d+\.\s*.+\?)|$)",
            r"([^\n]+\?)\s*(.+?)(?=(?:\n\s*\n)|(?:[^\n]+\?)|$)",
        ]

        for pattern in question_patterns:
            matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
            for q, a in matches:
                q = q.strip()
                a = re.sub(r"\s+", " ", a).strip()
                if q and a:
                    pairs.append({"question": q, "answer": a})
            if pairs:
                break

        return pairs

    def train_with_document(self, text: str) -> bool:
        """Fine-tune a small Hugging Face model on extracted Q/A pairs.

        This is CPU-friendly and intentionally minimal: it trains for a tiny
        number of steps and saves the updated model locally so future answers
        can use it.
        """
        qa_pairs = self._extract_qa_pairs(text)
        if not qa_pairs:
            print("No Q/A pairs detected in document; skipping fine-tuning.")
            return False

        qa_pairs = qa_pairs[:64]

        model_name = "google/flan-t5-small"
        output_dir = Path(__file__).parent.parent / "data" / "fine_tuned_models" / "answer_generator"
        output_dir.mkdir(parents=True, exist_ok=True)

        try:
            from transformers import (
                AutoModelForSeq2SeqLM,
                AutoTokenizer,
                DataCollatorForSeq2Seq,
                Seq2SeqTrainer,
                Seq2SeqTrainingArguments,
            )
        except Exception as e:
            print(f"Training unavailable (transformers import failed): {str(e)}")
            return False

        try:
            # Torch is optional for fine-tuning; ensure it's available before proceeding.
            try:
                import torch
            except Exception as e:
                print(f"Training unavailable (torch import failed): {str(e)}")
                return False

            tokenizer = AutoTokenizer.from_pretrained(model_name)
            model = AutoModelForSeq2SeqLM.from_pretrained(model_name)

            class _QADataset(torch.utils.data.Dataset):
                def __init__(self, tokenizer, pairs, max_source_length=384, max_target_length=128):
                    self.encodings = []
                    self.labels = []
                    for item in pairs:
                        q = item["question"]
                        a = item["answer"]
                        prompt = (
                            "You are an AI assistant answering questions using the user's personal knowledge base.\n\n"
                            f"Context:\n- {a}\n\nQuestion:\n{q}\n\nAnswer:"
                        )
                        enc = tokenizer(
                            prompt,
                            truncation=True,
                            padding="max_length",
                            max_length=max_source_length,
                            return_tensors="pt",
                        )
                        with tokenizer.as_target_tokenizer():
                            lab = tokenizer(
                                a,
                                truncation=True,
                                padding="max_length",
                                max_length=max_target_length,
                                return_tensors="pt",
                            )
                        self.encodings.append({k: v.squeeze(0) for k, v in enc.items()})
                        self.labels.append(lab["input_ids"].squeeze(0))

                def __len__(self):
                    return len(self.encodings)

                def __getitem__(self, idx):
                    item = dict(self.encodings[idx])
                    labels = self.labels[idx].clone()
                    labels[labels == tokenizer.pad_token_id] = -100
                    item["labels"] = labels
                    return item

            train_dataset = _QADataset(tokenizer, qa_pairs)
            data_collator = DataCollatorForSeq2Seq(tokenizer=tokenizer, model=model)

            args = Seq2SeqTrainingArguments(
                output_dir=str(output_dir),
                overwrite_output_dir=True,
                num_train_epochs=1,
                per_device_train_batch_size=2,
                learning_rate=5e-5,
                logging_steps=5,
                save_strategy="no",
                report_to=[],
                fp16=False,
                no_cuda=True,
            )

            trainer = Seq2SeqTrainer(
                model=model,
                args=args,
                train_dataset=train_dataset,
                data_collator=data_collator,
                tokenizer=tokenizer,
            )

            trainer.train()

            model.save_pretrained(output_dir)
            tokenizer.save_pretrained(output_dir)
            print(f"Fine-tuned model saved to: {output_dir}")
            return True
        except Exception as e:
            print(f"Error during fine-tuning: {str(e)}")
            return False

    def _process_qa_content(self, text):
        """Process text to extract question-answer pairs with better formatting."""
        if not text:
            return text

        # Common question patterns to look for
        question_patterns = [
            r'(?:Question|Q)[:\s]\s*(.+?)\s*(?:\n|$)(.+?)(?=(?:\n\s*\n)|(?:Question|Q)[:\s]|\Z)',
            r'(\d+\.\s*.+?\?)\s*(.+?)(?=(?:\n\s*\n)|(?:\d+\.\s*.+\?)|$)',
            r'([^\n]+\?)\s*(.+?)(?=(?:\n\s*\n)|(?:[^\n]+\?)|$)'
        ]

        for pattern in question_patterns:
            matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
            if matches:
                result = []
                for question, answer in matches:
                    question = question.strip()
                    answer = re.sub(r'\s+', ' ', answer).strip()
                    if question and answer:
                        result.append(f"## {question}\n\n{answer}\n")
                if result:
                    return "\n".join(result)

        return text
